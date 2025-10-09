import io

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocxConverter:
    """Converter for DOCX files to markdown format."""

    def __init__(self):
        self.image_handler = None
        self.extracted_images = {}  # Map: rId -> filename

    def convert(self, file, include_metadata=True, image_handler=None):
        """
        Convert DOCX file to markdown.

        Args:
            file: Streamlit uploaded file object
            include_metadata: Whether to include document metadata
            image_handler: Optional ImageHandler instance for extracting images

        Returns:
            str: Markdown content
        """
        try:
            self.image_handler = image_handler
            self.extracted_images = {}

            # Read the file content
            file_content = io.BytesIO(file.read())
            doc = Document(file_content)

            # Extract images if handler provided
            if self.image_handler:
                self._extract_all_images(doc)

            markdown_lines = []

            # Add metadata if requested
            if include_metadata:
                markdown_lines.extend(self._extract_metadata(doc, file.name))
                markdown_lines.append("")

            # Process paragraphs
            for paragraph in doc.paragraphs:
                markdown_line = self._convert_paragraph(paragraph)
                if markdown_line:
                    markdown_lines.append(markdown_line)

                # Check for inline images in the paragraph
                image_refs = self._extract_paragraph_images(paragraph)
                for image_ref in image_refs:
                    if image_ref in self.extracted_images:
                        markdown_lines.append(
                            f"\n![Image](assets/{self.extracted_images[image_ref]})\n"
                        )

            # Process tables
            for table in doc.tables:
                table_markdown = self._convert_table(table)
                if table_markdown:
                    markdown_lines.extend(table_markdown)
                    markdown_lines.append("")

            return "\n".join(markdown_lines)

        except Exception as e:
            raise Exception(f"Error converting DOCX file: {str(e)}")

    def _extract_metadata(self, doc, filename):
        """Extract document metadata."""
        metadata = [
            "---",
            f'title: "{filename}"',
            'source_format: "DOCX"',
        ]

        try:
            props = doc.core_properties
            if props.author:
                metadata.append(f'author: "{props.author}"')
            if props.created:
                metadata.append(f'created: "{props.created.isoformat()}"')
            if props.modified:
                metadata.append(f'modified: "{props.modified.isoformat()}"')
            if props.subject:
                metadata.append(f'subject: "{props.subject}"')
        except Exception:
            pass  # Skip metadata if not available

        metadata.append("---")
        return metadata

    def _convert_paragraph(self, paragraph):
        """Convert a paragraph to markdown."""
        if not paragraph.text.strip():
            return ""

        text = paragraph.text.strip()

        # Handle different paragraph styles
        style_name = paragraph.style.name.lower()

        if "heading" in style_name:
            # Extract heading level
            level = 1
            if "heading 1" in style_name:
                level = 1
            elif "heading 2" in style_name:
                level = 2
            elif "heading 3" in style_name:
                level = 3
            elif "heading 4" in style_name:
                level = 4
            elif "heading 5" in style_name:
                level = 5
            elif "heading 6" in style_name:
                level = 6

            return f"{'#' * level} {text}"

        # Handle alignment
        alignment = paragraph.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            text = f"<center>{text}</center>"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            text = f"<div align='right'>{text}</div>"

        # Process runs for inline formatting
        formatted_text = self._process_runs(paragraph.runs)

        return formatted_text if formatted_text.strip() else text

    def _process_runs(self, runs):
        """Process runs for inline formatting."""
        result = ""

        for run in runs:
            text = run.text

            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            if run.underline:
                text = f"<u>{text}</u>"

            result += text

        return result

    def _convert_table(self, table):
        """Convert a table to markdown format."""
        if not table.rows:
            return []

        markdown_table = []

        # Process header row
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]

        if any(header_cells):  # Only create table if there's content
            # Create header
            markdown_table.append("| " + " | ".join(header_cells) + " |")
            markdown_table.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

            # Process data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # Only add row if there's content
                    markdown_table.append("| " + " | ".join(cells) + " |")

        return markdown_table

    def _extract_all_images(self, doc):
        """Extract all images from the document."""
        if not self.image_handler:
            return  # No image handler, skip extraction

        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    r_id = rel.rId
                    image_data = image_part.blob

                    # Optimize and save the image
                    optimized_data, ext = self.image_handler.optimize_image(image_data)
                    filename = self.image_handler.save_image(
                        optimized_data, ext, prefix="docx_img"
                    )

                    # Store mapping
                    self.extracted_images[r_id] = filename

        except Exception as e:
            print(f"Warning: Could not extract images: {str(e)}")

    def _extract_paragraph_images(self, paragraph):
        """Extract image references from a paragraph."""
        image_refs = []

        try:
            # Look for drawing elements in the paragraph
            for run in paragraph.runs:
                # Check for inline images
                if "graphic" in run._element.xml:
                    # Parse the XML to find image references
                    # XML namespace URLs - cannot be shortened
                    ns_main = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"  # noqa: E501
                    ns_draw = "{http://schemas.openxmlformats.org/drawingml/2006/main}"  # noqa: E501
                    ns_rel = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"  # noqa: E501
                    for drawing in run._element.findall(f".//{ns_main}drawing"):
                        for blip in drawing.findall(f".//{ns_draw}blip"):
                            embed = blip.get(f"{ns_rel}embed")
                            if embed:
                                image_refs.append(embed)
        except Exception:
            pass  # Silently skip if image extraction fails

        return image_refs
