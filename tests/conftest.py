"""Pytest configuration and shared fixtures."""

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
import pytest
from docx import Document


@pytest.fixture
def temp_test_dir(tmp_path):
    """Create a temporary test directory."""
    test_dir = tmp_path / "test_files"
    test_dir.mkdir()
    return test_dir


@pytest.fixture
def sample_csv_content():
    """Sample CSV content for testing."""
    return """Name,Age,City,Country
John Doe,30,New York,USA
Jane Smith,25,London,UK
Bob Johnson,35,Toronto,Canada
Alice Brown,28,Sydney,Australia"""


@pytest.fixture
def sample_csv_file(sample_csv_content):
    """Create a mock CSV file object."""
    file_obj = io.BytesIO(sample_csv_content.encode("utf-8"))
    file_obj.name = "test_data.csv"
    file_obj.seek(0)
    return file_obj


@pytest.fixture
def sample_dataframe():
    """Create a sample pandas DataFrame."""
    return pd.DataFrame(
        {
            "Name": ["John Doe", "Jane Smith", "Bob Johnson", "Alice Brown"],
            "Age": [30, 25, 35, 28],
            "City": ["New York", "London", "Toronto", "Sydney"],
            "Country": ["USA", "UK", "Canada", "Australia"],
        }
    )


@pytest.fixture
def sample_txt_content():
    """Sample text content for testing."""
    return """# Test Document

This is a test document with multiple paragraphs.

## Section 1
This is the first section with some content.

## Section 2
This is the second section with more content.

- List item 1
- List item 2
- List item 3
"""


@pytest.fixture
def sample_txt_file(sample_txt_content):
    """Create a mock TXT file object."""
    file_obj = io.BytesIO(sample_txt_content.encode("utf-8"))
    file_obj.name = "test_document.txt"
    file_obj.seek(0)
    return file_obj


@pytest.fixture
def sample_docx_file(temp_test_dir):
    """Create a sample DOCX file."""
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Content of section 1.")
    doc.add_heading("Section 2", level=2)
    doc.add_paragraph("Content of section 2.")

    docx_path = temp_test_dir / "test_document.docx"
    doc.save(str(docx_path))

    # Create a file-like object
    with open(docx_path, "rb") as f:
        file_obj = io.BytesIO(f.read())
    file_obj.name = "test_document.docx"
    file_obj.seek(0)
    return file_obj


@pytest.fixture
def sample_wxr_content():
    """Sample WordPress WXR content for testing."""
    return """<?xml version="1.0" encoding="UTF-8"?>
<rss version="2.0"
    xmlns:excerpt="http://wordpress.org/export/1.2/excerpt/"
    xmlns:content="http://purl.org/rss/1.0/modules/content/"
    xmlns:wfw="http://wellformedweb.org/CommentAPI/"
    xmlns:dc="http://purl.org/dc/elements/1.1/"
    xmlns:wp="http://wordpress.org/export/1.2/">
<channel>
    <title>Test Blog</title>
    <item>
        <title>Test Post</title>
        <link>http://example.com/test-post</link>
        <pubDate>Mon, 01 Jan 2024 12:00:00 +0000</pubDate>
        <dc:creator><![CDATA[admin]]></dc:creator>
        <content:encoded><![CDATA[<p>This is test content.</p>]]></content:encoded>
        <excerpt:encoded><![CDATA[Test excerpt]]></excerpt:encoded>
        <wp:post_id>1</wp:post_id>
        <wp:post_date>2024-01-01 12:00:00</wp:post_date>
        <wp:post_type>post</wp:post_type>
        <wp:status>publish</wp:status>
        <category domain="category" nicename="test"><![CDATA[Test]]></category>
        <category domain="post_tag" nicename="example"><![CDATA[example]]></category>
    </item>
</channel>
</rss>"""


@pytest.fixture
def sample_wxr_file(sample_wxr_content):
    """Create a mock WXR file object."""
    file_obj = io.BytesIO(sample_wxr_content.encode("utf-8"))
    file_obj.name = "wordpress_export.wxr"
    file_obj.seek(0)
    return file_obj


@pytest.fixture
def sample_metadata():
    """Sample metadata dictionary for testing."""
    return {
        "title": "Test Document",
        "author": "Test Author",
        "date": "2024-01-01",
        "categories": ["Tech", "Tutorial"],
        "tags": ["python", "testing", "pytest"],
        "excerpt": "This is a test excerpt for the document.",
        "status": "publish",
    }


@pytest.fixture
def sample_markdown_with_frontmatter():
    """Sample markdown content with frontmatter."""
    return """---
title: "Test Document"
date: 2024-01-01
author: "Test Author"
tags:
  - python
  - testing
categories:
  - Tech
---

# Test Document

This is the content of the document.
"""


@pytest.fixture
def sample_html_content():
    """Sample HTML content for testing."""
    return """
    <html>
        <head><title>Test Page</title></head>
        <body>
            <h1>Main Heading</h1>
            <p>First paragraph with <strong>bold text</strong>.</p>
            <ul>
                <li>Item 1</li>
                <li>Item 2</li>
            </ul>
            <img src="image.jpg" alt="Test Image">
        </body>
    </html>
    """


@pytest.fixture
def mock_streamlit_file():
    """Create a mock Streamlit uploaded file."""

    class MockUploadedFile:
        def __init__(self, content: bytes, name: str):
            self._content = content
            self.name = name
            self.size = len(content)
            self._position = 0

        def read(self):
            self._position = len(self._content)
            return self._content

        def seek(self, position):
            self._position = position

        def getvalue(self):
            return self._content

    return MockUploadedFile


@pytest.fixture
def sample_seo_metadata():
    """Sample SEO metadata for testing."""
    return {
        "title": "Test Page | My Website",
        "description": "This is a test page description for SEO purposes.",
        "keywords": ["test", "seo", "metadata"],
        "og_title": "Test Page",
        "og_description": "Test description for Open Graph",
        "og_image": "/images/test.jpg",
        "twitter_card": "summary_large_image",
    }


@pytest.fixture
def temp_output_dir(tmp_path):
    """Create a temporary output directory."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    return output_dir


@pytest.fixture
def sample_image_data():
    """Sample base64 image data."""
    # 1x1 transparent PNG
    return "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
